"""Document processing tools for parsing and extraction."""

from typing import Any, Dict, List, Optional, BinaryIO
from pydantic import BaseModel, Field
from pathlib import Path
from datetime import datetime, timezone
import logging
import io

from app.tools.base import BaseTool, ToolMetadata, ToolCategory, ToolError


# ========== PDF Parser Tool ==========

class PDFParserInput(BaseModel):
    """PDF parser input model."""

    file_path: Optional[str] = Field(default=None, description="Path to PDF file")
    file_bytes: Optional[bytes] = Field(default=None, description="PDF file bytes")
    extract_text: bool = Field(default=True, description="Extract text content")
    extract_metadata: bool = Field(default=True, description="Extract PDF metadata")
    extract_images: bool = Field(default=False, description="Extract images")
    pages: Optional[List[int]] = Field(default=None, description="Specific pages to extract")


class PDFMetadata(BaseModel):
    """PDF metadata."""

    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    creator: Optional[str] = None
    producer: Optional[str] = None
    creation_date: Optional[datetime] = None
    modification_date: Optional[datetime] = None
    page_count: int = 0


class PDFPage(BaseModel):
    """PDF page content."""

    page_number: int
    text: str
    images: List[Dict[str, Any]] = Field(default_factory=list)


class PDFParserOutput(BaseModel):
    """PDF parser output model."""

    text: str = ""
    pages: List[PDFPage] = Field(default_factory=list)
    metadata: Optional[PDFMetadata] = None
    images: List[Dict[str, Any]] = Field(default_factory=list)


class PDFParserTool(BaseTool[PDFParserInput, PDFParserOutput]):
    """Tool for parsing PDF documents."""

    metadata = ToolMetadata(
        name="pdf_parser",
        description="Extract text and metadata from PDF documents",
        category=ToolCategory.DOCUMENT,
        version="1.0.0",
        tags=["pdf", "document", "parsing", "extraction"],
        timeout_seconds=120.0,
    )

    InputModel = PDFParserInput
    OutputModel = PDFParserOutput

    async def _execute(self, input_data: PDFParserInput) -> PDFParserOutput:
        """Execute PDF parsing.

        Args:
            input_data: Parsing parameters.

        Returns:
            PDFParserOutput: Parsed content.
        """
        try:
            import pypdf
        except ImportError:
            raise ToolError("pypdf library not installed. Install with: pip install pypdf")

        # Get PDF data
        if input_data.file_path:
            pdf_reader = pypdf.PdfReader(input_data.file_path)
        elif input_data.file_bytes:
            pdf_reader = pypdf.PdfReader(io.BytesIO(input_data.file_bytes))
        else:
            raise ToolError("Either file_path or file_bytes must be provided")

        output = PDFParserOutput()

        # Extract metadata
        if input_data.extract_metadata:
            output.metadata = self._extract_metadata(pdf_reader)

        # Extract text
        if input_data.extract_text:
            pages_to_process = input_data.pages or list(range(len(pdf_reader.pages)))
            all_text = []

            for page_num in pages_to_process:
                if page_num >= len(pdf_reader.pages):
                    continue

                page = pdf_reader.pages[page_num]
                text = page.extract_text() or ""
                all_text.append(text)

                output.pages.append(
                    PDFPage(
                        page_number=page_num + 1,
                        text=text,
                    )
                )

            output.text = "\n\n".join(all_text)

        # Extract images
        if input_data.extract_images:
            output.images = self._extract_images(pdf_reader)

        return output

    def _extract_metadata(self, pdf_reader: Any) -> PDFMetadata:
        """Extract PDF metadata.

        Args:
            pdf_reader: PyPDF reader object.

        Returns:
            PDFMetadata: Extracted metadata.
        """
        metadata = pdf_reader.metadata or {}

        def parse_date(date_str: Optional[str]) -> Optional[datetime]:
            if not date_str:
                return None
            try:
                # PDF date format: D:YYYYMMDDHHmmSS
                date_str = date_str.replace("D:", "")[:14]
                return datetime.strptime(date_str, "%Y%m%d%H%M%S")
            except Exception:
                return None

        return PDFMetadata(
            title=metadata.get("/Title"),
            author=metadata.get("/Author"),
            subject=metadata.get("/Subject"),
            creator=metadata.get("/Creator"),
            producer=metadata.get("/Producer"),
            creation_date=parse_date(metadata.get("/CreationDate")),
            modification_date=parse_date(metadata.get("/ModDate")),
            page_count=len(pdf_reader.pages),
        )

    def _extract_images(self, pdf_reader: Any) -> List[Dict[str, Any]]:
        """Extract images from PDF.

        Args:
            pdf_reader: PyPDF reader object.

        Returns:
            List[Dict]: Extracted images.
        """
        images = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                for image_key in page.images.keys():
                    image = page.images[image_key]
                    images.append(
                        {
                            "page": page_num + 1,
                            "name": image.name,
                            "data": image.data,
                            "image_type": image.image.type if hasattr(image.image, "type") else None,
                        }
                    )
            except Exception as e:
                self._logger.warning(f"Failed to extract images from page {page_num + 1}: {e}")
        return images


# ========== DOCX Parser Tool ==========

class DOCXParserInput(BaseModel):
    """DOCX parser input model."""

    file_path: Optional[str] = Field(default=None, description="Path to DOCX file")
    file_bytes: Optional[bytes] = Field(default=None, description="DOCX file bytes")
    extract_text: bool = Field(default=True, description="Extract text content")
    extract_metadata: bool = Field(default=True, description="Extract document metadata")
    extract_images: bool = Field(default=False, description="Extract embedded images")
    extract_tables: bool = Field(default=False, description="Extract tables")


class DOCXMetadata(BaseModel):
    """DOCX metadata."""

    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    description: Optional[str] = None
    created: Optional[datetime] = None
    modified: Optional[datetime] = None
    last_modified_by: Optional[str] = None
    revision: Optional[str] = None


class DOCXTable(BaseModel):
    """DOCX table."""

    rows: List[List[str]] = Field(default_factory=list)


class DOCXParserOutput(BaseModel):
    """DOCX parser output model."""

    text: str = ""
    metadata: Optional[DOCXMetadata] = None
    images: List[Dict[str, Any]] = Field(default_factory=list)
    tables: List[DOCXTable] = Field(default_factory=list)


class DOCXParserTool(BaseTool[DOCXParserInput, DOCXParserOutput]):
    """Tool for parsing Word documents."""

    metadata = ToolMetadata(
        name="docx_parser",
        description="Extract text and metadata from Word documents",
        category=ToolCategory.DOCUMENT,
        version="1.0.0",
        tags=["docx", "word", "document", "parsing"],
        timeout_seconds=60.0,
    )

    InputModel = DOCXParserInput
    OutputModel = DOCXParserOutput

    async def _execute(self, input_data: DOCXParserInput) -> DOCXParserOutput:
        """Execute DOCX parsing.

        Args:
            input_data: Parsing parameters.

        Returns:
            DOCXParserOutput: Parsed content.
        """
        try:
            from docx import Document
        except ImportError:
            raise ToolError("python-docx library not installed. Install with: pip install python-docx")

        # Get document
        if input_data.file_path:
            doc = Document(input_data.file_path)
        elif input_data.file_bytes:
            doc = Document(io.BytesIO(input_data.file_bytes))
        else:
            raise ToolError("Either file_path or file_bytes must be provided")

        output = DOCXParserOutput()

        # Extract text
        if input_data.extract_text:
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            output.text = "\n\n".join(paragraphs)

        # Extract metadata
        if input_data.extract_metadata:
            output.metadata = self._extract_metadata(doc)

        # Extract tables
        if input_data.extract_tables:
            output.tables = self._extract_tables(doc)

        # Extract images
        if input_data.extract_images:
            output.images = self._extract_images(doc)

        return output

    def _extract_metadata(self, doc: Any) -> DOCXMetadata:
        """Extract DOCX metadata.

        Args:
            doc: Document object.

        Returns:
            DOCXMetadata: Extracted metadata.
        """
        core_props = doc.core_properties
        return DOCXMetadata(
            title=core_props.title,
            author=core_props.author,
            subject=core_props.subject,
            keywords=core_props.keywords,
            description=core_props.description,
            created=core_props.created,
            modified=core_props.modified,
            last_modified_by=core_props.last_modified_by,
            revision=core_props.revision,
        )

    def _extract_tables(self, doc: Any) -> List[DOCXTable]:
        """Extract tables from document.

        Args:
            doc: Document object.

        Returns:
            List[DOCXTable]: Extracted tables.
        """
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(DOCXTable(rows=rows))
        return tables

    def _extract_images(self, doc: Any) -> List[Dict[str, Any]]:
        """Extract images from document.

        Args:
            doc: Document object.

        Returns:
            List[Dict]: Extracted images.
        """
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    images.append(
                        {
                            "name": rel.target_ref,
                            "data": rel.target_part.blob,
                        }
                    )
                except Exception as e:
                    self._logger.warning(f"Failed to extract image: {e}")
        return images


# ========== Markdown Parser Tool ==========

class MarkdownParserInput(BaseModel):
    """Markdown parser input model."""

    content: Optional[str] = Field(default=None, description="Markdown content string")
    file_path: Optional[str] = Field(default=None, description="Path to markdown file")
    extract_headings: bool = Field(default=True, description="Extract heading structure")
    extract_links: bool = Field(default=True, description="Extract all links")
    extract_code_blocks: bool = Field(default=True, description="Extract code blocks")
    convert_to_html: bool = Field(default=False, description="Convert to HTML")


class MarkdownHeading(BaseModel):
    """Markdown heading."""

    level: int
    text: str
    line_number: int


class MarkdownLink(BaseModel):
    """Markdown link."""

    text: str
    url: str
    line_number: int


class MarkdownCodeBlock(BaseModel):
    """Markdown code block."""

    language: Optional[str] = None
    code: str
    line_number: int


class MarkdownParserOutput(BaseModel):
    """Markdown parser output model."""

    text: str = ""
    html: Optional[str] = None
    headings: List[MarkdownHeading] = Field(default_factory=list)
    links: List[MarkdownLink] = Field(default_factory=list)
    code_blocks: List[MarkdownCodeBlock] = Field(default_factory=list)


class MarkdownParserTool(BaseTool[MarkdownParserInput, MarkdownParserOutput]):
    """Tool for parsing and rendering Markdown."""

    metadata = ToolMetadata(
        name="markdown_parser",
        description="Parse and render Markdown content",
        category=ToolCategory.DOCUMENT,
        version="1.0.0",
        tags=["markdown", "parsing", "rendering"],
        timeout_seconds=30.0,
    )

    InputModel = MarkdownParserInput
    OutputModel = MarkdownParserOutput

    async def _execute(self, input_data: MarkdownParserInput) -> MarkdownParserOutput:
        """Execute Markdown parsing.

        Args:
            input_data: Parsing parameters.

        Returns:
            MarkdownParserOutput: Parsed content.
        """
        try:
            import markdown
        except ImportError:
            raise ToolError("markdown library not installed. Install with: pip install markdown")

        # Get content
        if input_data.content:
            content = input_data.content
        elif input_data.file_path:
            with open(input_data.file_path, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            raise ToolError("Either content or file_path must be provided")

        output = MarkdownParserOutput(text=content)

        # Convert to HTML
        if input_data.convert_to_html:
            output.html = markdown.markdown(content, extensions=["extra", "codehilite"])

        # Extract headings
        if input_data.extract_headings:
            output.headings = self._extract_headings(content)

        # Extract links
        if input_data.extract_links:
            output.links = self._extract_links(content)

        # Extract code blocks
        if input_data.extract_code_blocks:
            output.code_blocks = self._extract_code_blocks(content)

        return output

    def _extract_headings(self, content: str) -> List[MarkdownHeading]:
        """Extract headings from markdown.

        Args:
            content: Markdown content.

        Returns:
            List[MarkdownHeading]: Extracted headings.
        """
        headings = []
        lines = content.split("\n")
        for line_num, line in enumerate(lines, 1):
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.strip("#").strip()
                headings.append(
                    MarkdownHeading(level=level, text=text, line_number=line_num)
                )
        return headings

    def _extract_links(self, content: str) -> List[MarkdownLink]:
        """Extract links from markdown.

        Args:
            content: Markdown content.

        Returns:
            List[MarkdownLink]: Extracted links.
        """
        import re

        links = []
        pattern = r"\[([^\]]+)\]\(([^)]+)\)"
        lines = content.split("\n")

        for line_num, line in enumerate(lines, 1):
            for match in re.finditer(pattern, line):
                links.append(
                    MarkdownLink(
                        text=match.group(1),
                        url=match.group(2),
                        line_number=line_num,
                    )
                )
        return links

    def _extract_code_blocks(self, content: str) -> List[MarkdownCodeBlock]:
        """Extract code blocks from markdown.

        Args:
            content: Markdown content.

        Returns:
            List[MarkdownCodeBlock]: Extracted code blocks.
        """
        blocks = []
        lines = content.split("\n")
        in_block = False
        block_start = 0
        block_lang = None
        block_lines = []

        for line_num, line in enumerate(lines, 1):
            if line.startswith("```"):
                if not in_block:
                    in_block = True
                    block_start = line_num
                    block_lang = line.strip("```").strip() or None
                    block_lines = []
                else:
                    in_block = False
                    blocks.append(
                        MarkdownCodeBlock(
                            language=block_lang,
                            code="\n".join(block_lines),
                            line_number=block_start,
                        )
                    )
            elif in_block:
                block_lines.append(line)

        return blocks


# ========== Table Extraction Tool ==========

class TableExtractionInput(BaseModel):
    """Table extraction input model."""

    file_path: Optional[str] = Field(default=None, description="Path to document")
    file_bytes: Optional[bytes] = Field(default=None, description="File bytes")
    file_type: str = Field(..., description="File type (pdf, docx, xlsx, html)")
    detection_method: str = Field(default="auto", description="Detection method")


class TableExtractionOutput(BaseModel):
    """Table extraction output model."""

    tables: List[Dict[str, Any]] = Field(default_factory=list)
    table_count: int = 0


class TableExtractionTool(BaseTool[TableExtractionInput, TableExtractionOutput]):
    """Tool for extracting tables from documents."""

    metadata = ToolMetadata(
        name="table_extractor",
        description="Extract tables from various document formats",
        category=ToolCategory.DOCUMENT,
        version="1.0.0",
        tags=["table", "extraction", "document"],
        timeout_seconds=120.0,
    )

    InputModel = TableExtractionInput
    OutputModel = TableExtractionOutput

    async def _execute(self, input_data: TableExtractionInput) -> TableExtractionOutput:
        """Execute table extraction.

        Args:
            input_data: Extraction parameters.

        Returns:
            TableExtractionOutput: Extracted tables.
        """
        file_type = input_data.file_type.lower()

        if file_type == "pdf":
            return await self._extract_from_pdf(input_data)
        elif file_type == "docx":
            return await self._extract_from_docx(input_data)
        elif file_type == "xlsx":
            return await self._extract_from_xlsx(input_data)
        elif file_type == "html":
            return await self._extract_from_html(input_data)
        else:
            raise ToolError(f"Unsupported file type: {file_type}")

    async def _extract_from_pdf(self, input_data: TableExtractionInput) -> TableExtractionOutput:
        """Extract tables from PDF.

        Args:
            input_data: Extraction parameters.

        Returns:
            TableExtractionOutput: Extracted tables.
        """
        try:
            import tabula
        except ImportError:
            raise ToolError("tabula-py not installed. Install with: pip install tabula-py")

        if input_data.file_path:
            tables = tabula.read_pdf(input_data.file_path, pages="all", multiple_tables=True)
        elif input_data.file_bytes:
            tables = tabula.read_pdf(
                io.BytesIO(input_data.file_bytes),
                pages="all",
                multiple_tables=True,
            )
        else:
            raise ToolError("Either file_path or file_bytes must be provided")

        return TableExtractionOutput(
            tables=[table.to_dict() for table in tables],
            table_count=len(tables),
        )

    async def _extract_from_docx(self, input_data: TableExtractionInput) -> TableExtractionOutput:
        """Extract tables from DOCX.

        Args:
            input_data: Extraction parameters.

        Returns:
            TableExtractionOutput: Extracted tables.
        """
        try:
            from docx import Document
        except ImportError:
            raise ToolError("python-docx not installed")

        if input_data.file_path:
            doc = Document(input_data.file_path)
        elif input_data.file_bytes:
            doc = Document(io.BytesIO(input_data.file_bytes))
        else:
            raise ToolError("Either file_path or file_bytes must be provided")

        tables = []
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append({"rows": rows})

        return TableExtractionOutput(tables=tables, table_count=len(tables))

    async def _extract_from_xlsx(self, input_data: TableExtractionInput) -> TableExtractionOutput:
        """Extract tables from XLSX.

        Args:
            input_data: Extraction parameters.

        Returns:
            TableExtractionOutput: Extracted tables.
        """
        try:
            import pandas as pd
        except ImportError:
            raise ToolError("pandas not installed")

        if input_data.file_path:
            xls = pd.ExcelFile(input_data.file_path)
        elif input_data.file_bytes:
            xls = pd.ExcelFile(io.BytesIO(input_data.file_bytes))
        else:
            raise ToolError("Either file_path or file_bytes must be provided")

        tables = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            tables.append(
                {
                    "sheet": sheet_name,
                    "data": df.to_dict(),
                    "columns": list(df.columns),
                }
            )

        return TableExtractionOutput(tables=tables, table_count=len(tables))

    async def _extract_from_html(self, input_data: TableExtractionInput) -> TableExtractionOutput:
        """Extract tables from HTML.

        Args:
            input_data: Extraction parameters.

        Returns:
            TableExtractionOutput: Extracted tables.
        """
        try:
            import pandas as pd
        except ImportError:
            raise ToolError("pandas not installed")

        if input_data.file_path:
            with open(input_data.file_path, "r", encoding="utf-8") as f:
                html_content = f.read()
        elif input_data.file_bytes:
            html_content = input_data.file_bytes.decode("utf-8")
        else:
            raise ToolError("Either file_path or file_bytes must be provided")

        tables_list = pd.read_html(html_content)

        return TableExtractionOutput(
            tables=[table.to_dict() for table in tables_list],
            table_count=len(tables_list),
        )


# ========== Image OCR Tool ==========

class ImageOCRInput(BaseModel):
    """Image OCR input model."""

    file_path: Optional[str] = Field(default=None, description="Path to image file")
    file_bytes: Optional[bytes] = Field(default=None, description="Image file bytes")
    language: str = Field(default="eng", description="OCR language code")
    extract_layout: bool = Field(default=False, description="Extract layout information")


class ImageOCROutput(BaseModel):
    """Image OCR output model."""

    text: str = ""
    confidence: Optional[float] = None
    language: str = "eng"
    layout: Optional[Dict[str, Any]] = None


class ImageOCRTool(BaseTool[ImageOCRInput, ImageOCROutput]):
    """Tool for Optical Character Recognition on images."""

    metadata = ToolMetadata(
        name="image_ocr",
        description="Extract text from images using OCR",
        category=ToolCategory.DOCUMENT,
        version="1.0.0",
        tags=["ocr", "image", "text-extraction"],
        timeout_seconds=60.0,
    )

    InputModel = ImageOCRInput
    OutputModel = ImageOCROutput

    async def _execute(self, input_data: ImageOCRInput) -> ImageOCROutput:
        """Execute OCR on image.

        Args:
            input_data: OCR parameters.

        Returns:
            ImageOCROutput: Extracted text.
        """
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            raise ToolError("pytesseract or PIL not installed")

        # Get image
        if input_data.file_path:
            image = Image.open(input_data.file_path)
        elif input_data.file_bytes:
            image = Image.open(io.BytesIO(input_data.file_bytes))
        else:
            raise ToolError("Either file_path or file_bytes must be provided")

        # Perform OCR
        config = f"--oem 3 --psm 6"
        if input_data.extract_layout:
            config += " -c tessedit_create_hocr=1"

        ocr_data = pytesseract.image_to_data(
            image,
            lang=input_data.language,
            config=config,
            output_type=pytesseract.Output.DICT,
        )

        # Extract text
        text = pytesseract.image_to_string(
            image,
            lang=input_data.language,
            config=config,
        )

        # Calculate average confidence
        confidences = [int(c) for c in ocr_data["conf"] if int(c) != -1]
        avg_confidence = sum(confidences) / len(confidences) if confidences else None

        output = ImageOCROutput(
            text=text.strip(),
            confidence=avg_confidence,
            language=input_data.language,
        )

        # Extract layout if requested
        if input_data.extract_layout:
            output.layout = ocr_data

        return output
